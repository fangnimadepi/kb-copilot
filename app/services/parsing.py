"""文档解析：PDF / Word / Markdown -> 带页码的文本单元序列。

统一抽象是 TextUnit（一段文本 + 页码）：
- PDF：每页一个 unit，页码从 1 起（引用溯源的基础）
- docx：Word 没有稳定的"页"概念（分页由渲染决定），按段落输出，page=0
- Markdown：按行输出保留标题结构，page=0

表格处理：PDF 用 PyMuPDF 的 find_tables 把表格重排为 Markdown 表格文本，
避免默认文本抽取把单元格拍平成无结构的字串——这是年报场景的关键差异点。
"""

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class TextUnit:
    text: str
    page: int  # 1 起；无页概念的格式为 0


class UnsupportedFormat(Exception):
    pass


def parse_file(path: str | Path) -> list[TextUnit]:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix in (".md", ".markdown", ".txt"):
        return _parse_markdown(path)
    raise UnsupportedFormat(f"不支持的格式: {suffix}（支持 pdf/docx/md/txt）")


def _parse_pdf(path: Path) -> list[TextUnit]:
    units: list[TextUnit] = []
    with fitz.open(path) as doc:
        for page_no, page in enumerate(doc, start=1):
            # 先抽表格并记录其区域，正文抽取时跳过这些区域避免重复
            tables = page.find_tables()
            table_bboxes = [t.bbox for t in tables.tables]
            table_texts = [_table_to_markdown(t) for t in tables.tables]

            body = _page_text_excluding(page, table_bboxes)
            parts = [body, *[t for t in table_texts if t]]
            text = "\n\n".join(p for p in parts if p.strip())
            if text.strip():
                units.append(TextUnit(text=text, page=page_no))
    return units


def _page_text_excluding(page: "fitz.Page", bboxes: list) -> str:
    if not bboxes:
        return page.get_text("text")
    blocks = page.get_text("blocks")
    kept = []
    for b in blocks:
        rect = fitz.Rect(b[:4])
        if not any(rect.intersects(fitz.Rect(bb)) for bb in bboxes):
            kept.append(b[4])
    return "\n".join(kept)


def _table_to_markdown(table: "fitz.table.Table") -> str:
    rows = table.extract()
    if not rows:
        return ""
    lines = []
    header = [(_cell(c)) for c in rows[0]]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "---|" * len(header))
    for row in rows[1:]:
        lines.append("| " + " | ".join(_cell(c) for c in row) + " |")
    return "\n".join(lines)


def _cell(value) -> str:
    return (value or "").replace("\n", " ").strip()


def _parse_docx(path: Path) -> list[TextUnit]:
    doc = DocxDocument(str(path))
    units: list[TextUnit] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 保留标题层级信息，结构化分块策略会用到
        if para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "")
            prefix = "#" * int(level) if level.isdigit() else "#"
            text = f"{prefix} {text}"
        units.append(TextUnit(text=text, page=0))
    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if rows:
            lines = ["| " + " | ".join(rows[0]) + " |", "|" + "---|" * len(rows[0])]
            lines += ["| " + " | ".join(r) + " |" for r in rows[1:]]
            units.append(TextUnit(text="\n".join(lines), page=0))
    return units


def _parse_markdown(path: Path) -> list[TextUnit]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [TextUnit(text=line, page=0) for line in text.splitlines() if line.strip()]
